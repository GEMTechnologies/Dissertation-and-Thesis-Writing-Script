from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

def create_word_document(file_name, content):
    document = Document()

    # Set the page size to A4
    section = document.sections[0]
    section.page_height = Pt(842)
    section.page_width = Pt(595)
    section.start_type

    # Set font size and style for the document
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Add chapter one headings to the document
    add_chapter_one_headings(document)

    # Add content to the document
    paragraphs = content.split("\n")
    for paragraph in paragraphs:
        # Add paragraph
        p = document.add_paragraph()
        p.add_run(paragraph)

    # Save the document with the specified file name
    document.save(file_name)

def add_chapter_one_headings(document):
    headings = [
        ("CHAPTER ONE: INTRODUCTION", 0),
        ("1.0 Introduction", 1),
        ("1.1 Background of the study", 1),
        # ... [all other headings]
        ("1.14 Chapter One conclusion", 1)
    ]

    for text, level in headings:
        p = document.add_heading(text, level=level)
        if level == 0:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.runs[0]
            run.bold = True
            font = run.font
            font.color.rgb = (0, 0, 0)  # Black color
            font.size = Pt(14)
            font.name = 'Times New Roman'
            paragraph_format = p.paragraph_format
            paragraph_format.space_after = Pt(12)
        elif level == 1:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.runs[0]
            run.bold = True
            font = run.font
            font.color.rgb = (0, 0, 0)  # Black color
            font.size = Pt(12)
            font.name = 'Times New Roman'
            paragraph_format = p.paragraph_format
            paragraph_format.space_after = Pt(12)
        # ... [You can add more formatting for other heading levels if needed]

    # Add a page break after the headings
    document.add_page_break()
